from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUT = ROOT / "Joshua_Nehohwa_CV.docx"
PDF_OUT = ROOT / "public" / "Joshua_Nehohwa_CV.pdf"

CONTACT = [
    "Cape Town, South Africa",
    "nehohwajoshua@gmail.com",
    "+27 63 857 3965",
    "linkedin.com/in/joshua-nehohwa-b4b97b229",
    "github.com/jnehohwa",
]

SUMMARY = (
    "Software Engineering student and junior developer with experience across "
    "React/Next.js, TypeScript, Python, PHP/MySQL, Supabase, cloud fundamentals, "
    "and AI data evaluation. Strong academic performer with an 86% average, "
    "chairperson of Vossie DevClub, and practical experience leading technical "
    "workshops, building product prototypes, and working remotely with detailed "
    "AI annotation guidelines."
)

SECTIONS = [
    (
        "Education",
        [
            (
                "BSc Information Technology (Software Engineering), Eduvos University",
                "Mowbray Campus, Cape Town | Expected graduation: Nov 2026 | Academic average: 86%",
                [
                    "Relevant modules: Data Structures and Algorithms, Software Architecture, Mobile App Development, Network Security, Cloud-Based Technologies, Mathematics 1A/1B, and AI Ethics.",
                ],
            )
        ],
    ),
    (
        "Technical Skills",
        [
            (
                "Languages",
                "Python, JavaScript, TypeScript, Visual Basic, Java, PHP",
                [],
            ),
            (
                "Frameworks and tools",
                "React, Next.js, Node.js, Tailwind CSS, Supabase, Clerk Auth, Git, Linux, Docker basics, Figma, Justinmind",
                [],
            ),
            (
                "Databases and cloud",
                "PostgreSQL, MySQL, Supabase, AWS EC2, S3, IAM, Lambda",
                [],
            ),
        ],
    ),
    (
        "Leadership and Projects",
        [
            (
                "Chairperson, Vossie DevClub",
                "2025 - Present",
                [
                    "Lead technical workshops, student developer programs, community outreach, and campus innovation events.",
                    "Directed school outreach sessions teaching high school students foundational programming concepts.",
                    "Organize hackathons, mentorship initiatives, and student-focused developer activities.",
                ],
            ),
            (
                "HackJam Innovation Platform, Team Leader",
                "2025",
                [
                    "Led a team that designed and implemented a campus-wide innovation platform for idea submission, voting, and mentor feedback.",
                    "Built the frontend with React, Next.js, TypeScript, Tailwind CSS, and component libraries.",
                    "Added gamification concepts including points, badges, mentor picks, and placeholders for AI-based scoring and feedback.",
                    "Drove user research, prototyping, and final pitching; the project placed 4th in the university-wide HackJam competition.",
                ],
            ),
            (
                "Innosimm ERP System, Developer",
                "2025 - Ongoing",
                [
                    "Designing a custom ERP system for a vehicle and tyre business covering inventory, sales, purchasing, and basic finance tracking.",
                    "Implementing a modular Next.js architecture with Supabase and Clerk for authentication and data persistence.",
                    "Modeling real workflows such as stock movement, quotations, and invoices to replace manual spreadsheets and WhatsApp-based processes.",
                ],
            ),
            (
                "KasiSwap Marketplace",
                "2026",
                [
                    "Built marketplace surfaces across React/TypeScript and PHP/MySQL implementations, including listings, authentication-aware workflows, checkout-style order states, messaging, disputes, reviews, and admin moderation.",
                    "Created rubric-aligned documentation, evidence packs, deployment notes, and code samples for an academic deliverable.",
                ],
            ),
        ],
    ),
    (
        "Professional Experience",
        [
            (
                "Freelance AI Data Specialist, Data Annotations AI",
                "Remote | Jan 2024 - Nov 2024",
                [
                    "Refined large language model behavior by transforming raw HTML and text data into high-quality annotations.",
                    "Provided detailed feedback on model responses and edge cases for production-level AI systems used in Google and Microsoft contract work.",
                    "Worked independently with complex guidelines and evolving policies while meeting quality and throughput targets.",
                    "Built practical understanding of prompts, context windows, evaluation metrics, and real-world AI behavior.",
                ],
            ),
            (
                "Lift Attendant, Snowshoe Mountain",
                "West Virginia, USA | 2024-2025; returning 2025-2026",
                [
                    "Recognized multiple times as a top-performing employee for reliability and guest service.",
                    "Managed guest safety, lift operations, and high-pressure decisions in challenging weather conditions.",
                    "Strengthened resilience, teamwork, and cross-cultural communication in a diverse international environment.",
                ],
            ),
        ],
    ),
    (
        "Achievements and Certifications",
        [
            (
                "Golden Key International Honour Society",
                "Top 15% of university cohort, 2025",
                [],
            ),
            (
                "AWS Certifications",
                "AWS Cloud Practitioner and AWS Solutions Architect Associate",
                [],
            ),
            (
                "IELTS Academic",
                "Overall Band 8.5 / CEFR C2: Listening 8.5, Reading 9.0, Writing 6.5, Speaking 9.0",
                [],
            ),
        ],
    ),
    (
        "Interests",
        [
            (
                "Technical and personal interests",
                "AI research, cybersecurity, cloud infrastructure, software systems design, trail running, basketball, golf, tennis, snowboarding, and chess.",
                [],
            )
        ],
    ),
]


def set_run_style(run, size=10.5, bold=False, color="1F2937"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text.upper())
    set_run_style(run, size=10.5, bold=True, color="0F766E")


def add_docx_role(document, title, meta, bullets):
    title_line = document.add_paragraph()
    title_line.paragraph_format.space_after = Pt(1)
    title_run = title_line.add_run(title)
    set_run_style(title_run, size=10.5, bold=True, color="111827")

    if meta:
        meta_run = title_line.add_run(f" | {meta}")
        set_run_style(meta_run, size=9.4, color="4B5563")

    for bullet in bullets:
        item = document.add_paragraph(style="List Bullet")
        item.paragraph_format.space_after = Pt(1)
        run = item.add_run(bullet)
        set_run_style(run, size=9.4, color="374151")


def build_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Joshua Nehohwa")
    set_run_style(title_run, size=22, bold=True, color="111827")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle_run = subtitle.add_run(
        "Software Engineering Student | Junior Software Developer"
    )
    set_run_style(subtitle_run, size=10.5, bold=True, color="0F766E")

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(10)
    contact_run = contact.add_run(" | ".join(CONTACT))
    set_run_style(contact_run, size=8.6, color="4B5563")

    add_docx_heading(document, "Professional Summary")
    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(4)
    summary_run = summary.add_run(SUMMARY)
    set_run_style(summary_run, size=9.7, color="374151")

    for section_title, entries in SECTIONS:
        add_docx_heading(document, section_title)
        for title_text, meta, bullets in entries:
            add_docx_role(document, title_text, meta, bullets)

    document.save(DOCX_OUT)


def build_pdf():
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=colors.HexColor("#111827"),
        alignment=TA_CENTER,
        spaceAfter=3,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=colors.HexColor("#0F766E"),
        alignment=TA_CENTER,
        spaceAfter=5,
    )
    contact_style = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.2,
        leading=10,
        textColor=colors.HexColor("#4B5563"),
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.2,
        leading=13,
        textColor=colors.HexColor("#0F766E"),
        spaceBefore=9,
        spaceAfter=4,
    )
    role_style = ParagraphStyle(
        "Role",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9.8,
        leading=12,
        textColor=colors.HexColor("#111827"),
        spaceAfter=1,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.1,
        leading=12.2,
        textColor=colors.HexColor("#374151"),
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=14,
        firstLineIndent=-8,
        spaceAfter=1.2,
    )

    story.extend(
        [
            Paragraph("Joshua Nehohwa", title_style),
            Paragraph(
                "Software Engineering Student | Junior Software Developer",
                subtitle_style,
            ),
            Paragraph(" | ".join(CONTACT), contact_style),
            Paragraph("PROFESSIONAL SUMMARY", heading_style),
            Paragraph(SUMMARY, body_style),
        ]
    )

    for section_title, entries in SECTIONS:
        if section_title == "Professional Experience":
            story.append(PageBreak())

        story.append(Paragraph(section_title.upper(), heading_style))
        for title_text, meta, bullets in entries:
            role = f"<b>{title_text}</b>"
            if meta:
                role += f" | <font color='#4B5563'>{meta}</font>"
            story.append(Paragraph(role, role_style))
            if bullets:
                for bullet in bullets:
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{bullet}", bullet_style))
            story.append(Spacer(1, 1.5))

    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
        title="Joshua Nehohwa CV",
        author="Joshua Nehohwa",
    )
    document.build(story)


def main():
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
